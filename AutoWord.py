import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def estilo(run, size=16, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold

def agregar_linea(doc, etiqueta, valor, espacio_antes=0):
    p = doc.add_paragraph()

    run1 = p.add_run(etiqueta)
    estilo(run1, bold=True)

    run2 = p.add_run(valor)
    estilo(run2)

    p.paragraph_format.space_before = Pt(espacio_antes)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return p

# Crear Documento
doc = Document()

# Ruta base del script
base_dir = os.path.dirname(os.path.abspath(__file__))

# === ENCABEZADO ===
section = doc.sections[0]
header = section.header

# === HEADER ===
header_paragraph = header.paragraphs[0]
run = header_paragraph.add_run()

ruta_imagen = os.path.join(base_dir, "encabezado.jpg") # Nombre del archivo encabezado (.JPG-PNG)
run.add_picture(ruta_imagen, width=Inches(2.94))
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_paragraph.paragraph_format.left_indent = 0
header_paragraph.paragraph_format.space_before = 0
header_paragraph.paragraph_format.space_after = 0

# ===== PRIMERA PÁGINA =====
# === Título ===
titulo = doc.add_heading("Título del Documento", 0)
titulo.paragraph_format.space_before = Pt(150)
run = titulo.runs[0]
run.font.name = "Arial"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Tamaño y color del titulo
run.font.size = Pt(52)
run.font.color.rgb = RGBColor(0, 0, 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Borrar linea bajo el titulo
style = doc.styles['Heading 1']
pPr = style._element.get_or_add_pPr()
pPr = titulo._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'nil')
pBdr.append(bottom)
pPr.append(pBdr)

# === Subtítulo ===
parrafo = doc.add_paragraph()
run = parrafo.add_run("Descripción Genérica (Opcional)")
estilo(run)
parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# === Texto Estudiantes ===
agregar_linea(doc, "Estudiantes: ", "Integrante1 - Integrante2 - Integrante3", 230)

# === Texto Profesor(a) ===
agregar_linea(doc, "Profesor(a): ", "Nombre Profesor(a)")

# === Texto Asignatura ===
agregar_linea(doc, "Asignatura: ", "Asignatura Solicitada")

# === Texto Sección ===
agregar_linea(doc, "Sección: ", "RQY1102-008D")

# ========= SEGUNDA PÁGINA =========
texto2 = doc.add_paragraph()
run2 = texto2.add_run("Índice")
estilo(run2)
run2.bold = True
texto2.alignment = WD_ALIGN_PARAGRAPH.LEFT
texto2.paragraph_format.space_before = Pt(20)

# ========= TERCERA PÁGINA =========
doc.add_page_break()
texto3 = doc.add_paragraph()
run3 = texto3.add_run("1. Introducción")
estilo(run3)
run3.bold = True
texto3.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ========= CUARTA PÁGINA =========
doc.add_page_break()
texto4 = doc.add_paragraph()
run4 = texto4.add_run("2. Descripción General")
estilo(run4)
run4.bold = True
texto4.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ========= QUINTA PÁGINA =========
doc.add_page_break()
texto5 = doc.add_paragraph()
run5 = texto5.add_run("3. Propuesta de Planificación")
estilo(run5)
run5.bold = True
texto5.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ========= SÉPTIMA PÁGINA =========
doc.add_page_break()
texto6 = doc.add_paragraph()
run6 = texto6.add_run("4. Conclusión")
estilo(run6)
run6.bold = True
texto6.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Guardar documento
doc.save(os.path.join(base_dir, "Documento.docx"))